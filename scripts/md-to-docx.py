#!/usr/bin/env python3
"""Convert the strategy Markdown docs into Word (.docx) files.

No pandoc needed — uses python-docx. Handles headings, tables, bullet/numbered
lists, blockquotes, code fences, horizontal rules and inline **bold**/*italic*/
`code`/[links]. Brand-ish styling: navy serif headings, clean body.

Usage:
  python3 scripts/md-to-docx.py                 # all *.md under "05 Strategie"
  python3 scripts/md-to-docx.py path/to/file.md ...   # specific files

Each .docx is written next to its .md.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
STRAT = ROOT.parent / "05 Strategie"

NAVY = RGBColor(0x12, 0x19, 0x20)
CLAY = RGBColor(0xB9, 0x82, 0x72)
GREY = RGBColor(0x33, 0x41, 0x4E)
HEAD_FONT = "Georgia"
BODY_FONT = "Calibri"

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|_[^_]+?_|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))")


def add_runs(p, text):
    """Parse inline markdown in `text` and append styled runs to paragraph p."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif (tok.startswith("*") and tok.endswith("*")) or (tok.startswith("_") and tok.endswith("_")):
            r = p.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("[") and "](" in tok:
            label = tok[1:tok.index("]")]
            p.add_run(label)
        else:
            p.add_run(tok)
    return p


def style_body(p):
    for r in p.runs:
        if not r.font.name:
            r.font.name = BODY_FONT


def convert(md_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = BODY_FONT
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    in_code = False
    code_buf = []

    def flush_code():
        if not code_buf:
            return
        p = doc.add_paragraph()
        r = p.add_run("\n".join(code_buf))
        r.font.name = "Consolas"; r.font.size = Pt(9.5); r.font.color.rgb = GREY
        code_buf.clear()

    while i < n:
        line = lines[i]

        # code fences
        if line.strip().startswith("```"):
            if in_code:
                flush_code(); in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line); i += 1; continue

        # tables: a block of lines starting with |
        if line.lstrip().startswith("|") and i + 1 < n and re.match(r"\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            tbl_lines = []
            while i < n and lines[i].lstrip().startswith("|"):
                tbl_lines.append(lines[i]); i += 1
            rows = []
            for ti, tl in enumerate(tbl_lines):
                if ti == 1:
                    continue  # separator
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                rows.append(cells)
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=0, cols=cols)
                table.style = "Light Grid Accent 1"
                for ri, cells in enumerate(rows):
                    cells = cells + [""] * (cols - len(cells))
                    wr = table.add_row().cells
                    for ci, ctext in enumerate(cells):
                        cp = wr[ci].paragraphs[0]
                        add_runs(cp, ctext)
                        for run in cp.runs:
                            run.font.size = Pt(9.5)
                            if ri == 0:
                                run.bold = True
                doc.add_paragraph()
            continue

        # horizontal rule
        if re.match(r"^\s*(---|\*\*\*|___)\s*$", line):
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_paragraph()
            add_runs(p, text)
            sz = {1: 20, 2: 15, 3: 12.5}.get(level, 11)
            for r in p.runs:
                r.font.name = HEAD_FONT; r.bold = True
                r.font.size = Pt(sz)
                r.font.color.rgb = NAVY if level > 1 else CLAY
            p.space_before = Pt(10); p.space_after = Pt(4)
            i += 1
            continue

        # blockquote (collect consecutive > lines)
        if line.lstrip().startswith(">"):
            buf = []
            while i < n and lines[i].lstrip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", lines[i])); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, " ".join(x for x in buf if x.strip()))
            for r in p.runs:
                r.italic = True; r.font.color.rgb = GREY
            continue

        # bullet list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            style = "List Bullet" if indent < 2 else "List Bullet 2"
            p = doc.add_paragraph(style=style)
            add_runs(p, m.group(2)); style_body(p)
            i += 1
            continue

        # numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2)); style_body(p)
            i += 1
            continue

        # blank line
        if not line.strip():
            i += 1
            continue

        # normal paragraph
        p = doc.add_paragraph()
        add_runs(p, line.strip()); style_body(p)
        i += 1

    flush_code()
    out = md_path.with_suffix(".docx")
    doc.save(out)
    return out


def main():
    args = sys.argv[1:]
    if args:
        targets = [Path(a) for a in args]
    else:
        targets = sorted(STRAT.rglob("*.md"))
    for md in targets:
        out = convert(md)
        try:
            shown = out.resolve().relative_to(STRAT.parent.resolve())
        except ValueError:
            shown = out
        print(f"  {shown}")
    print(f"{len(targets)} Word-Dokumente erstellt.")


if __name__ == "__main__":
    main()
